"""
General utilities for the statejobs-helper project.
"""

import io
import os
import re

from docx import Document
from PyPDF2 import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# pdfkit optional
try:
    import pdfkit

    PDFKIT_AVAILABLE = True
except (ImportError, OSError):
    # pdfkit is not installed (need to add to requirements.txt)
    PDFKIT_AVAILABLE = False

# Liberation Sans is installed via Dockerfile and is a suitable replacement for Arial/Helvetica
DEFAULT_FONT_FACE = "Liberation Sans"
DEFAULT_CSS_FILE = "static/css/html_to_pdf.css"


def fill_template(file, data):
    """Replace placeholders like {{ name }} in a text or HTML template."""
    if hasattr(file, "read"):
        template_text = file.read().decode("utf-8", errors="ignore")
    else:
        template_text = str(file)

    def replacer(match):
        key = match.group(1).strip()
        return str(data.get(key, f"{{{{ {key} }}}}"))

    return re.sub(r"\{\{\s*(.*?)\s*\}\}", replacer, template_text)


def _convert_text_to_html(text_content: str) -> str:
    """
    Helper function to convert raw text content (split into header and body)
    into formatted HTML paragraphs.

    This function was created to reduce the complexity and local variable count
    of extract_text_and_html (Pylint R0914, R0912, R0915).
    """
    header_marker = "---END HEADER---"
    if header_marker in text_content:
        parts = text_content.split(header_marker, 1)
        header_raw_block = parts[0]
        body_raw_block = parts[1]
    else:
        parts = text_content.split("\n\n", 1)
        header_raw_block = parts[0]
        body_raw_block = parts[1] if len(parts) > 1 else ""

    # Header HTML Generation
    header_chunks = header_raw_block.strip().split("\n\n")
    header_html_parts = []
    for chunk in header_chunks:
        chunk = chunk.strip()
        if chunk:
            lines = [line.strip() for line in chunk.split("\n") if line.strip()]
            header_html_parts.append("<p>" + " <br> ".join(lines) + "</p>")
        else:
            header_html_parts.append("<p><br></p>")
    header_html = "".join(header_html_parts)

    # Body HTML Generation
    body_raw_block = body_raw_block.strip()
    body_paragraphs_and_breaks = body_raw_block.split("\n\n")
    # Guaranteed break after header/greeting
    body_html_parts = ["<p><br></p>"]

    for block in body_paragraphs_and_breaks:
        block_content = block.strip()
        if block_content:
            body_html_parts.append(f"<p>{block_content.replace('\n', ' ')}</p>")
            body_html_parts.append("<p><br></p>")

    if body_html_parts and body_html_parts[-1] == "<p><br></p>":
        body_html_parts.pop()

    body_html = "".join(body_html_parts).strip()
    return header_html + body_html


def extract_text_and_html(file_storage):
    """
    Extract text, HTML, and detected font size from uploaded template files.
    Returns: (text_content, html_content, detected_font_size)
    """
    filename = file_storage.filename.lower()

    file_bytes = file_storage.read()
    file_storage.seek(0)

    def normalize_text(s: str) -> str:
        s = s.replace("\r\n", "\n").replace("\r", "\n")
        s = s.replace("\xa0", " ").replace("\t", " ")
        return re.sub(r"\n{3,}", "\n\n", s.strip())

    text_content = ""

    # Default font size if detection fails or file is not DOCX
    detected_font_size = "12pt"

    if filename.endswith(".txt"):
        text_content = normalize_text(file_bytes.decode("utf-8", errors="ignore"))

    elif filename.endswith(".docx"):
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [normalize_text(p.text) for p in doc.paragraphs]
        text_content = "\n\n".join(paragraphs)

        for p in doc.paragraphs:
            if p.text.strip():
                try:
                    # python-docx stores size in half points
                    if (
                        p.style.font.size is not None
                        and p.style.font.size.pt is not None
                    ):
                        detected_font_size = f"{p.style.font.size.pt}pt"
                        break

                    if (
                        p.runs
                        and p.runs[0].font.size is not None
                        and p.runs[0].font.size.pt is not None
                    ):
                        detected_font_size = f"{p.runs[0].font.size.pt}pt"
                        break
                except (
                    AttributeError,
                    ValueError,
                ):
                    pass

    elif filename.endswith(".pdf"):
        pdf = PdfReader(io.BytesIO(file_bytes))
        text_pages = [normalize_text(page.extract_text() or "") for page in pdf.pages]
        text_content = "\n\n".join(text_pages)

    else:
        raise ValueError(f"Unsupported file type: {filename}")

    html_content = _convert_text_to_html(text_content)

    return text_content, html_content, detected_font_size


def text_to_pdf(text, font_size="12pt"):
    """
    Generates a PDF using ReportLab paragraph handling.
    Now accepts font_size to respect styling from the caller.
    """
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)

    # Extract the numeric font size (in points)
    try:
        size_pt = float(font_size.replace("pt", "").strip())
        if size_pt <= 0:
            size_pt = 12.0
    except ValueError:
        size_pt = 12.0

    _, height = letter

    y = height - 72
    # Adjust line height to be proportional to the new font size (size + 2 points)
    line_height = size_pt + 2

    for paragraph in text.split("\n\n"):

        text_object = c.beginText(72, y)
        # Use the requested default font
        text_object.setFont(DEFAULT_FONT_FACE, size_pt)

        lines = [line.strip() for line in paragraph.splitlines() if line.strip()]

        if lines:
            for line in lines:
                text_object.textLine(line)
                y -= line_height
                if y < 72:
                    c.drawText(text_object)
                    c.showPage()
                    y = height - 72
                    text_object = c.beginText(72, y)
                    text_object.setFont(DEFAULT_FONT_FACE, size_pt)

            c.drawText(text_object)
        else:
            y -= line_height

        y -= line_height

        if y < 72:
            c.showPage()
            y = height - 72

    c.save()
    buffer.seek(0)
    return buffer


def html_to_pdf(html_content, font_size="12pt"):
    """
    Generates a PDF from HTML using pdfkit (wkhtmltopdf), or falls back to text_to_pdf.
    """
    if PDFKIT_AVAILABLE:
        try:
            DEFAULT_CSS = f"""
                     body, body * {{
                        font-family: "{DEFAULT_FONT_FACE}", sans-serif !important;
                        line-height: 1.2 !important;
                        margin: 0;
                        padding: 0;
                    }}
                    p {{
                        margin-top: 0;
                        margin-bottom: 0;
                        text-indent: 0;
                    }}
                    /* Tweak for line breaks */
                    p br {{
                        line-height: 0.8;
                        display: block;
                        content: "";
                        margin-bottom: -0.2em;
                    }}

                    {{dynamic_css}}
            """
            FONT_SIZE_CSS = f"""
                    body, body * {{
                        font-size: {font_size} !important;
                    }}
            """

            final_css = None
            css_template_source = DEFAULT_CSS

            try:
                absolute_path = os.path.abspath(DEFAULT_CSS_FILE)
                with open(absolute_path, "r") as f:
                    css_template_source = f.read()

            except Exception as e:
                pass

            escaped_css_template = css_template_source.replace("{", "{{").replace(
                "}", "}}"
            )
            escaped_css_template = escaped_css_template.replace(
                "{{dynamic_css}}", "{dynamic_css}"
            )
            final_css = escaped_css_template.format(dynamic_css=FONT_SIZE_CSS)

            print(f"Generated final_css: {final_css}")

            # Define wkhtmltopdf options
            options = {
                "page-size": "Letter",
                "margin-top": "1in",
                "margin-right": "1in",
                "margin-bottom": "1in",
                "margin-left": "1in",
                "encoding": "UTF-8",
                # This option helps with Docker/headless environments
                "quiet": "",
            }

            # Embed the CSS directly into the HTML payload
            html_content_inlined = html_content.strip()
            styled_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    {final_css}
                </style>
            </head>
            <body>
                {html_content_inlined}
            </body>
            </html>
            """

            print(f"----------------\n{styled_html}\n----------------\n")
            # 3. Generate PDF to bytes
            # pdfkit.from_string(html, output_path, options). False as output_path returns bytes.
            pdf_bytes = pdfkit.from_string(styled_html, False, options=options)

            pdf_buffer = io.BytesIO(pdf_bytes)
            pdf_buffer.seek(0)
            return pdf_buffer

        except Exception as e:
            # You might want to log this exception for debugging in your final app
            print(f"pdfkit failed, falling back to ReportLab: {e}")
            pass  # Silently fail to the fallback

    # Fallback to text_to_pdf
    text_content = html_content
    # The following lines convert the HTML back into simple text for ReportLab
    text_content = re.sub(r"</?p[^>]*>", "\n\n", text_content)
    text_content = text_content.replace("<br>", "\n")
    text_content = re.sub(r"<[^>]+>", "", text_content).strip()
    text_content = re.sub(r"\n{3,}", "\n\n", text_content)

    return text_to_pdf(text_content, font_size)
